"""
DocSummarizer GUI
Modern cross-platform GUI using CustomTkinter.
"""

import os
import threading
from pathlib import Path
from typing import Optional
import customtkinter as ctk
from tkinter import filedialog, messagebox

from document_parser import extract_text, get_document_info, SUPPORTED_EXTENSIONS
from model_manager import (
    Summarizer, download_model, is_model_downloaded,
    get_model_path, DEFAULT_MODEL, SMALL_MODEL
)


# Appearance settings
ctk.set_appearance_mode("System")  # "System", "Dark", "Light"
ctk.set_default_color_theme("blue")


class DocSummarizerApp(ctk.CTk):
    """Main application window."""

    def __init__(self):
        super().__init__()

        self.title("DocSummarizer - Offline Document Summarization")
        self.geometry("900x700")
        self.minsize(700, 500)

        self.summarizer: Optional[Summarizer] = None
        self.current_file: Optional[str] = None
        self.extracted_text: Optional[str] = None
        self._is_closing = False

        # Handle window close properly
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        self._create_widgets()
        self._check_model_status()

    def _on_close(self):
        """Clean shutdown when window is closed."""
        self._is_closing = True
        # Release model from memory
        if self.summarizer:
            del self.summarizer
            self.summarizer = None
        self.destroy()

    def _create_widgets(self):
        """Create all GUI widgets."""
        # Configure grid
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)

        # === Header Frame ===
        self.header_frame = ctk.CTkFrame(self)
        self.header_frame.grid(row=0, column=0, padx=10, pady=(10, 5), sticky="ew")
        self.header_frame.grid_columnconfigure(1, weight=1)

        self.title_label = ctk.CTkLabel(
            self.header_frame,
            text="DocSummarizer",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        self.title_label.grid(row=0, column=0, padx=10, pady=5)

        self.status_label = ctk.CTkLabel(
            self.header_frame,
            text="Checking model status...",
            font=ctk.CTkFont(size=12)
        )
        self.status_label.grid(row=0, column=1, padx=10, pady=5, sticky="e")

        # === File Selection Frame ===
        self.file_frame = ctk.CTkFrame(self)
        self.file_frame.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        self.file_frame.grid_columnconfigure(1, weight=1)

        self.select_btn = ctk.CTkButton(
            self.file_frame,
            text="Select File",
            command=self._select_file,
            width=120
        )
        self.select_btn.grid(row=0, column=0, padx=10, pady=10)

        self.file_label = ctk.CTkLabel(
            self.file_frame,
            text="No file selected",
            anchor="w"
        )
        self.file_label.grid(row=0, column=1, padx=10, pady=10, sticky="ew")

        self.select_folder_btn = ctk.CTkButton(
            self.file_frame,
            text="Batch (Folder)",
            command=self._select_folder,
            width=120
        )
        self.select_folder_btn.grid(row=0, column=2, padx=10, pady=10)

        # === Main Content Frame (Tabview) ===
        self.tabview = ctk.CTkTabview(self)
        self.tabview.grid(row=2, column=0, padx=10, pady=5, sticky="nsew")

        self.tab_summary = self.tabview.add("Summary")
        self.tab_text = self.tabview.add("Extracted Text")
        self.tab_settings = self.tabview.add("Settings")

        # Configure tabs
        self.tab_summary.grid_columnconfigure(0, weight=1)
        self.tab_summary.grid_rowconfigure(0, weight=1)
        self.tab_text.grid_columnconfigure(0, weight=1)
        self.tab_text.grid_rowconfigure(0, weight=1)

        # Summary tab
        self.summary_text = ctk.CTkTextbox(self.tab_summary, wrap="word")
        self.summary_text.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
        self.summary_text.insert("1.0", "Summary will appear here after processing...")
        self.summary_text.configure(state="disabled")

        # Extracted text tab
        self.extracted_textbox = ctk.CTkTextbox(self.tab_text, wrap="word")
        self.extracted_textbox.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")
        self.extracted_textbox.insert("1.0", "Extracted text will appear here...")
        self.extracted_textbox.configure(state="disabled")

        # Settings tab
        self._create_settings_tab()

        # === Controls Frame ===
        self.controls_frame = ctk.CTkFrame(self)
        self.controls_frame.grid(row=3, column=0, padx=10, pady=(5, 10), sticky="ew")
        self.controls_frame.grid_columnconfigure(1, weight=1)

        # Summary type selection
        self.summary_type_label = ctk.CTkLabel(self.controls_frame, text="Summary Type:")
        self.summary_type_label.grid(row=0, column=0, padx=(10, 5), pady=10)

        self.summary_type_var = ctk.StringVar(value="detailed")
        self.summary_type_menu = ctk.CTkOptionMenu(
            self.controls_frame,
            variable=self.summary_type_var,
            values=["brief", "detailed", "structured"],
            width=120
        )
        self.summary_type_menu.grid(row=0, column=1, padx=5, pady=10, sticky="w")

        # Progress bar
        self.progress_bar = ctk.CTkProgressBar(self.controls_frame)
        self.progress_bar.grid(row=0, column=2, padx=10, pady=10, sticky="ew")
        self.progress_bar.set(0)
        self.controls_frame.grid_columnconfigure(2, weight=1)

        # Action buttons
        self.summarize_btn = ctk.CTkButton(
            self.controls_frame,
            text="Summarize",
            command=self._start_summarization,
            width=120,
            state="disabled"
        )
        self.summarize_btn.grid(row=0, column=3, padx=5, pady=10)

        self.save_btn = ctk.CTkButton(
            self.controls_frame,
            text="Save Summary",
            command=self._save_summary,
            width=120,
            state="disabled"
        )
        self.save_btn.grid(row=0, column=4, padx=(5, 10), pady=10)

    def _create_settings_tab(self):
        """Create the settings tab content."""
        self.tab_settings.grid_columnconfigure(0, weight=1)

        # Model section
        model_label = ctk.CTkLabel(
            self.tab_settings,
            text="Model Settings",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        model_label.grid(row=0, column=0, padx=10, pady=(10, 5), sticky="w")

        self.model_info_label = ctk.CTkLabel(
            self.tab_settings,
            text="Model: Checking...",
            anchor="w"
        )
        self.model_info_label.grid(row=1, column=0, padx=10, pady=5, sticky="w")

        self.download_btn = ctk.CTkButton(
            self.tab_settings,
            text="Download Model",
            command=self._start_model_download,
            width=150
        )
        self.download_btn.grid(row=2, column=0, padx=10, pady=10, sticky="w")

        # Appearance section
        appearance_label = ctk.CTkLabel(
            self.tab_settings,
            text="Appearance",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        appearance_label.grid(row=3, column=0, padx=10, pady=(20, 5), sticky="w")

        self.appearance_var = ctk.StringVar(value="System")
        self.appearance_menu = ctk.CTkOptionMenu(
            self.tab_settings,
            variable=self.appearance_var,
            values=["System", "Light", "Dark"],
            command=self._change_appearance,
            width=150
        )
        self.appearance_menu.grid(row=4, column=0, padx=10, pady=5, sticky="w")

        # About section
        about_label = ctk.CTkLabel(
            self.tab_settings,
            text="About",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        about_label.grid(row=5, column=0, padx=10, pady=(20, 5), sticky="w")

        about_text = ctk.CTkLabel(
            self.tab_settings,
            text="DocSummarizer - Offline Document Summarization Tool\n"
                 "Uses a local AI model to summarize documents.\n"
                 "No internet required after initial model download.\n\n"
                 "Supported formats: PDF, DOCX, DOC, RTF, TXT, MD",
            anchor="w",
            justify="left"
        )
        about_text.grid(row=6, column=0, padx=10, pady=5, sticky="w")

    def _check_model_status(self):
        """Check if the model is downloaded and update UI accordingly."""
        if is_model_downloaded():
            self.status_label.configure(text="Model ready", text_color="green")
            self.model_info_label.configure(
                text=f"Model: {DEFAULT_MODEL['name']} ({DEFAULT_MODEL['size_gb']} GB) - Downloaded"
            )
            self.download_btn.configure(state="disabled", text="Model Downloaded")
            self._load_model()
        else:
            self.status_label.configure(text="Model not downloaded", text_color="orange")
            self.model_info_label.configure(
                text=f"Model: {DEFAULT_MODEL['name']} ({DEFAULT_MODEL['size_gb']} GB) - Not downloaded"
            )
            self.download_btn.configure(state="normal")

    def _load_model(self):
        """Load the model in a background thread."""
        def load():
            try:
                self.status_label.configure(text="Loading model...")
                model_path = get_model_path()
                self.summarizer = Summarizer(model_path)
                self.status_label.configure(text="Ready", text_color="green")
                self._update_button_states()
            except Exception as e:
                self.status_label.configure(text=f"Error loading model: {str(e)}", text_color="red")

        thread = threading.Thread(target=load, daemon=True)
        thread.start()

    def _start_model_download(self):
        """Start downloading the model in a background thread."""
        def download():
            self.download_btn.configure(state="disabled", text="Downloading...")
            self.progress_bar.set(0)

            def progress_callback(percent, message):
                self.progress_bar.set(percent / 100)
                self.status_label.configure(text=message)

            path, error = download_model(progress_callback=progress_callback)

            if error:
                self.status_label.configure(text=error, text_color="red")
                self.download_btn.configure(state="normal", text="Retry Download")
            else:
                self.download_btn.configure(text="Model Downloaded")
                self._load_model()

        thread = threading.Thread(target=download, daemon=True)
        thread.start()

    def _select_file(self):
        """Open file dialog to select a document."""
        filetypes = [
            ("All Supported", "*.pdf *.docx *.doc *.rtf *.txt *.md"),
            ("PDF Files", "*.pdf"),
            ("Word Documents", "*.docx *.doc"),
            ("Text Files", "*.txt *.md *.rtf"),
        ]

        filepath = filedialog.askopenfilename(
            title="Select Document",
            filetypes=filetypes
        )

        if filepath:
            self._process_file(filepath)

    def _select_folder(self):
        """Open folder dialog for batch processing."""
        folder = filedialog.askdirectory(title="Select Folder with Documents")

        if folder:
            # Find all supported files
            extensions = ('.pdf', '.docx', '.doc', '.rtf', '.txt', '.md')
            files = [f for f in Path(folder).iterdir()
                     if f.is_file() and f.suffix.lower() in extensions]

            if not files:
                messagebox.showinfo("No Files", "No supported documents found in the selected folder.")
                return

            result = messagebox.askyesno(
                "Batch Processing",
                f"Found {len(files)} document(s). Process all?"
            )

            if result:
                self._batch_process(files)

    def _process_file(self, filepath: str):
        """Extract text from the selected file."""
        self.current_file = filepath
        info = get_document_info(filepath)
        self.file_label.configure(text=f"{info['name']} ({info['size_mb']} MB)")

        # Extract text
        self.status_label.configure(text="Extracting text...")
        text, error = extract_text(filepath)

        if error:
            self.status_label.configure(text=error, text_color="red")
            self.extracted_text = None
        else:
            self.extracted_text = text
            self.status_label.configure(text="Text extracted", text_color="green")

            # Update extracted text view
            self.extracted_textbox.configure(state="normal")
            self.extracted_textbox.delete("1.0", "end")
            self.extracted_textbox.insert("1.0", text)
            self.extracted_textbox.configure(state="disabled")

        self._update_button_states()

    def _update_button_states(self):
        """Update button states based on current state."""
        can_summarize = (
            self.summarizer is not None and
            self.extracted_text is not None
        )
        self.summarize_btn.configure(state="normal" if can_summarize else "disabled")

    def _start_summarization(self):
        """Start the summarization process in a background thread."""
        if not self.summarizer or not self.extracted_text:
            return

        def summarize():
            self.summarize_btn.configure(state="disabled", text="Processing...")
            self.progress_bar.set(0.3)
            self.status_label.configure(text="Generating summary...")

            try:
                summary = self.summarizer.summarize(
                    self.extracted_text,
                    summary_type=self.summary_type_var.get()
                )

                self.progress_bar.set(1.0)
                self.status_label.configure(text="Summary complete", text_color="green")

                # Update summary text
                self.summary_text.configure(state="normal")
                self.summary_text.delete("1.0", "end")
                self.summary_text.insert("1.0", summary)
                self.summary_text.configure(state="disabled")

                self.save_btn.configure(state="normal")
                self.tabview.set("Summary")

            except Exception as e:
                self.status_label.configure(text=f"Error: {str(e)}", text_color="red")

            finally:
                self.summarize_btn.configure(state="normal", text="Summarize")
                self.progress_bar.set(0)

        thread = threading.Thread(target=summarize, daemon=True)
        thread.start()

    def _batch_process(self, files: list):
        """Process multiple files and save summaries."""
        output_folder = filedialog.askdirectory(title="Select Output Folder for Summaries")
        if not output_folder:
            return

        def process_batch():
            self.summarize_btn.configure(state="disabled")
            total = len(files)

            for i, filepath in enumerate(files):
                self.status_label.configure(text=f"Processing {i+1}/{total}: {filepath.name}")
                self.progress_bar.set((i + 0.5) / total)

                # Extract text
                text, error = extract_text(str(filepath))
                if error:
                    continue

                # Generate summary
                try:
                    summary = self.summarizer.summarize(
                        text,
                        summary_type=self.summary_type_var.get()
                    )

                    # Save summary
                    output_path = Path(output_folder) / f"{filepath.stem}_summary.txt"
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(f"Summary of: {filepath.name}\n")
                        f.write("=" * 50 + "\n\n")
                        f.write(summary)

                except Exception as e:
                    print(f"Error processing {filepath}: {e}")

                self.progress_bar.set((i + 1) / total)

            self.status_label.configure(text=f"Batch complete: {total} files processed", text_color="green")
            self.summarize_btn.configure(state="normal")
            messagebox.showinfo("Complete", f"Processed {total} files.\nSummaries saved to: {output_folder}")

        thread = threading.Thread(target=process_batch, daemon=True)
        thread.start()

    def _save_summary(self):
        """Save the current summary to a file."""
        if not self.current_file:
            return

        default_name = Path(self.current_file).stem + "_summary.txt"

        filepath = filedialog.asksaveasfilename(
            title="Save Summary",
            defaultextension=".txt",
            initialfile=default_name,
            filetypes=[
                ("Text File", "*.txt"),
                ("Markdown", "*.md"),
                ("Word Document", "*.docx"),
            ]
        )

        if filepath:
            self.summary_text.configure(state="normal")
            content = self.summary_text.get("1.0", "end-1c")
            self.summary_text.configure(state="disabled")

            if filepath.endswith('.docx'):
                from docx import Document
                doc = Document()
                doc.add_heading(f"Summary: {Path(self.current_file).name}", 0)
                doc.add_paragraph(content)
                doc.save(filepath)
            else:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)

            self.status_label.configure(text=f"Saved: {Path(filepath).name}", text_color="green")

    def _change_appearance(self, mode: str):
        """Change the application appearance mode."""
        ctk.set_appearance_mode(mode)


def main():
    """Main entry point."""
    app = DocSummarizerApp()
    app.mainloop()


if __name__ == "__main__":
    main()
