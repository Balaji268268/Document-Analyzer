"""
I/O helpers shared between the GUI and CLI.

Both entry points need to write summaries to text or docx files with the
same header format. Keeping that here means the format only has to be
maintained in one place.
"""

from pathlib import Path


def write_summary_txt(
    output_path: str | Path,
    *,
    source_name: str,
    summary: str,
    summary_type: str | None = None,
    separator_width: int = 50,
) -> None:
    """Write `summary` to a UTF-8 text file with a standard header.

    The header is:

        Summary of: <source_name>
        Type: <summary_type>          # omitted when summary_type is None
        ==================================================

        <summary body>
    """
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(f"Summary of: {source_name}\n")
        if summary_type:
            f.write(f"Type: {summary_type}\n")
        f.write("=" * separator_width + "\n\n")
        f.write(summary)


def write_summary_docx(
    output_path: str | Path,
    *,
    source_name: str,
    summary: str,
) -> None:
    """Write `summary` to a `.docx` file with `source_name` as the heading."""
    # Imported lazily; python-docx is only needed when the user opts into
    # docx output, and the CLI shouldn't pay for the import.
    from docx import Document

    doc = Document()
    doc.add_heading(f"Summary: {source_name}", 0)
    doc.add_paragraph(summary)
    doc.save(str(output_path))
