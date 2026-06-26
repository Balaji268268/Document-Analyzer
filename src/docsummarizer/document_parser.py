"""
Document Parser Module
Extracts text from various document formats: PDF, DOCX, TXT, RTF.

Legacy binary `.doc` (OLE compound) is intentionally not supported —
python-docx only reads the XML-based `.docx` format. The single
source-of-truth for "what we can parse" is `SUPPORTED_EXTENSION_SET`
below; dialog filters (`SUPPORTED_EXTENSIONS`), the extractor dispatch
in `extract_text`, and CLI directory scans all derive from it.
"""

from collections.abc import Callable
from pathlib import Path
from typing import TypedDict

import chardet


class DocumentInfo(TypedDict):
    """Basic metadata about a document file."""

    name: str
    extension: str
    size_bytes: int
    size_mb: float


def extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n\n".join(text_parts)


def extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(file_path)
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


def extract_from_rtf(file_path: str) -> str:
    """Extract text from an RTF file."""
    from striprtf.striprtf import rtf_to_text

    with Path(file_path).open(encoding="utf-8", errors="ignore") as f:
        rtf_content = f.read()

    return str(rtf_to_text(rtf_content))


def extract_from_txt(file_path: str) -> str:
    """Extract text from a plain text file with encoding detection."""
    raw_data = Path(file_path).read_bytes()

    detected = chardet.detect(raw_data)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    return raw_data.decode(encoding, errors="replace")


_EXTRACTORS: dict[str, Callable[[str], str]] = {
    ".pdf": extract_from_pdf,
    ".docx": extract_from_docx,
    ".rtf": extract_from_rtf,
    ".txt": extract_from_txt,
    ".md": extract_from_txt,
    ".text": extract_from_txt,
}

# Canonical set of supported extensions. Single source of truth — the
# dispatch above, the dialog filters below, and `find_documents` all
# read from it.
SUPPORTED_EXTENSION_SET: frozenset[str] = frozenset(_EXTRACTORS.keys())


def extract_text(file_path: str) -> tuple[str, str | None]:
    """Extract text from a document file.

    Returns:
        Tuple of (extracted_text, error_message). error_message is None on
        success; extracted_text is empty on failure.
    """
    path = Path(file_path)

    if not path.exists():
        return "", f"File not found: {file_path}"

    extractor = _EXTRACTORS.get(path.suffix.lower())
    if extractor is None:
        return "", f"Unsupported file format: {path.suffix.lower()}"

    try:
        text = extractor(file_path)
    except Exception as e:
        return "", f"Error extracting text: {e!s}"

    if not text.strip():
        return "", "No text could be extracted from the document"
    return text, None


def get_document_info(file_path: str) -> DocumentInfo:
    """Get basic information about a document."""
    path = Path(file_path)

    # Single stat() call; the previous code stat()'d twice and could race
    # on network filesystems (FileNotFoundError on the second call).
    try:
        size_bytes = path.stat().st_size
    except FileNotFoundError:
        size_bytes = 0

    return {
        "name": path.name,
        "extension": path.suffix.lower(),
        "size_bytes": size_bytes,
        "size_mb": round(size_bytes / (1024 * 1024), 2),
    }


def find_documents(path: Path) -> list[Path]:
    """Find all supported documents at `path`.

    If `path` is a single file, returns it (or empty if unsupported).
    If `path` is a directory, returns its supported files (non-recursive).
    """
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_EXTENSION_SET else []
    return [
        f for f in path.iterdir() if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSION_SET
    ]


class DocumentStats(TypedDict):
    """Rich metadata for the Summary/Extract header strip.

    ``pages`` is only meaningful for PDFs (``None`` otherwise); ``encoding`` is
    only detected for plain-text formats (``None`` otherwise). ``words``/``chars``
    are derived from the extracted text; ``parser`` names the engine used.
    """

    name: str
    pages: int | None
    words: int
    chars: int
    encoding: str | None
    parser: str


# Which extraction engine handles each extension — surfaced in the header strip.
_PARSER_LABELS: dict[str, str] = {
    ".pdf": "pypdf",
    ".docx": "python-docx",
    ".rtf": "striprtf",
    ".txt": "chardet",
    ".md": "chardet",
    ".text": "chardet",
}

# Extensions whose encoding we can meaningfully detect (chardet on raw bytes).
_TEXT_EXTENSIONS = frozenset({".txt", ".md", ".text"})


def _pdf_page_count(file_path: str) -> int | None:
    """Page count for a PDF, or ``None`` if it can't be read.

    A second, lightweight open (the extractor discards the page count); never
    raises, so a corrupt/missing PDF just yields ``None`` pages in the strip.
    """
    try:
        from pypdf import PdfReader

        return len(PdfReader(file_path).pages)
    except Exception:  # best-effort metadata; a bad PDF just yields None pages
        return None


def _detect_encoding(file_path: str) -> str | None:
    """Detected text encoding for a plain-text file, or ``None`` on read error."""
    try:
        raw = Path(file_path).read_bytes()
    except OSError:
        return None
    return chardet.detect(raw).get("encoding") or None


def analyze_document(file_path: str, text: str | None = None) -> DocumentStats:
    """Gather header-strip metadata for a document.

    ``words``/``chars`` come from ``text`` when supplied (the caller usually
    already extracted it), otherwise the text is extracted here. ``pages`` is
    PDF-only and ``encoding`` is plain-text-only; both degrade to ``None``.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    if text is None:
        text, _ = extract_text(file_path)
    return {
        "name": path.name,
        "pages": _pdf_page_count(file_path) if ext == ".pdf" else None,
        "words": len(text.split()),
        "chars": len(text),
        "encoding": _detect_encoding(file_path) if ext in _TEXT_EXTENSIONS else None,
        "parser": _PARSER_LABELS.get(ext, "unknown"),
    }


# Dialog filter list, derived from SUPPORTED_EXTENSION_SET so it can't drift.
SUPPORTED_EXTENSIONS = [
    ("All Supported", "*.pdf *.docx *.rtf *.txt *.md"),
    ("PDF Files", "*.pdf"),
    ("Word Documents", "*.docx"),
    ("RTF Files", "*.rtf"),
    ("Text Files", "*.txt *.md"),
]
