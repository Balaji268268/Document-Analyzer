"""
Document Parser Module
Extracts text from various document formats: PDF, DOCX, TXT, RTF
"""

import os
from pathlib import Path
from typing import Optional
import chardet


def extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n\n".join(text_parts)


def extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


def extract_from_rtf(file_path: str) -> str:
    """Extract text from an RTF file."""
    from striprtf.striprtf import rtf_to_text

    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        rtf_content = f.read()

    return rtf_to_text(rtf_content)


def extract_from_txt(file_path: str) -> str:
    """Extract text from a plain text file with encoding detection."""
    with open(file_path, 'rb') as f:
        raw_data = f.read()

    detected = chardet.detect(raw_data)
    encoding = detected.get('encoding', 'utf-8') or 'utf-8'

    return raw_data.decode(encoding, errors='replace')


def extract_text(file_path: str) -> tuple[str, Optional[str]]:
    """
    Extract text from a document file.

    Args:
        file_path: Path to the document

    Returns:
        Tuple of (extracted_text, error_message)
        If successful, error_message is None
    """
    path = Path(file_path)

    if not path.exists():
        return "", f"File not found: {file_path}"

    suffix = path.suffix.lower()

    extractors = {
        '.pdf': extract_from_pdf,
        '.docx': extract_from_docx,
        '.doc': extract_from_docx,  # May work for some .doc files
        '.rtf': extract_from_rtf,
        '.txt': extract_from_txt,
        '.md': extract_from_txt,
        '.text': extract_from_txt,
    }

    extractor = extractors.get(suffix)

    if extractor is None:
        return "", f"Unsupported file format: {suffix}"

    try:
        text = extractor(file_path)
        if not text.strip():
            return "", "No text could be extracted from the document"
        return text, None
    except Exception as e:
        return "", f"Error extracting text: {str(e)}"


def get_document_info(file_path: str) -> dict:
    """Get basic information about a document."""
    path = Path(file_path)

    return {
        'name': path.name,
        'extension': path.suffix.lower(),
        'size_bytes': path.stat().st_size if path.exists() else 0,
        'size_mb': round(path.stat().st_size / (1024 * 1024), 2) if path.exists() else 0,
    }


# Supported extensions for file dialogs
SUPPORTED_EXTENSIONS = [
    ("All Supported", "*.pdf *.docx *.doc *.rtf *.txt *.md"),
    ("PDF Files", "*.pdf"),
    ("Word Documents", "*.docx *.doc"),
    ("RTF Files", "*.rtf"),
    ("Text Files", "*.txt *.md"),
]
